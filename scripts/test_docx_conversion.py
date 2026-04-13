#!/usr/bin/env python3
"""
PDF → DOCX Conversion Prototype — Unit 4 Block B gate.

Standalone proof-of-concept that takes a single CAE release-notes PDF and
produces a Flightscape-templated DOCX. NOT wired into the pipeline. Run it
manually, eyeball the output in Word with Alex, record the verdict in the
PR description, and decide go/no-go before Unit 5 starts.

See PLAN_DOCS_PIPELINE.md §2 Block B and the plan file at
~/.claude/plans/linked-roaming-gray.md for the full design rationale.

Usage:

    python scripts/test_docx_conversion.py \\
        --pdf "patches/ACARS_V8_1/8.1.12.0/release_notes/8.1.12.0 - Release Notes.pdf" \\
        --output "docs_example/conversion_prototype/8.1.12.0.docx" \\
        --verbose

Optional:
    --template <path>         Override the Flightscape template path
    --mode {fast,hybrid}      Extraction backend (default: fast)
    --product "Operations Communication Manager"   Cover-page product name
    --no-cache                Re-run extraction even if cached JSON exists
    --json <path>             Skip extraction and use this pre-extracted JSON
                              (must point at an opendataloader-pdf JSON; the
                              <stem>_images/ folder must sit next to it). Useful
                              when Java isn't installed locally — exercise the
                              conversion logic against the cached V8.0 outputs
                              under docs_example/pdf_examples/8.0/extracted/fast/.

Extraction modes:

  - fast (default): opendataloader-pdf local Java extractor. No network.
  - hybrid:         opendataloader-pdf via the Docling Fast Server (IBM
                    Docling, layout-aware ML). Requires the local server:
                        opendataloader-pdf-hybrid --port 5002
                    Slower than fast but understands semantic structure
                    (heading vs body, image-to-paragraph anchoring, etc.).
                    If the server isn't reachable the script falls back to
                    fast with a clear log line.
"""

from __future__ import annotations

import argparse
import hashlib
import json
import logging
import os
import re
import shutil
import socket
import sys
from datetime import datetime
from pathlib import Path
from typing import Any, Iterable

# python-docx — installed in the venv but not yet in backend/requirements.txt
# (Unit 5 will productionise this).
from docx import Document
from docx.oxml.ns import qn
from docx.shared import Inches

# opendataloader-pdf is the only extractor we use in the prototype. Both
# 'fast' and 'hybrid-claude' modes route through this single library.
import opendataloader_pdf


# ─────────────────────────── constants ───────────────────────────

PROJECT_ROOT = Path(__file__).resolve().parent.parent
DEFAULT_TEMPLATE = (
    PROJECT_ROOT
    / "docs_example"
    / "templates"
    / "Flightscape-English-External Business Document.docx"
)
CACHE_ROOT = PROJECT_ROOT / "docs_example" / "conversion_prototype" / ".cache"
CLAUDE_CACHE_ROOT = CACHE_ROOT / "claude"

HYBRID_HOST = "127.0.0.1"
HYBRID_PORT = 5002

# Map opendataloader-pdf 'heading level' (1..N) → Flightscape style names.
# The template tops out at "Heading 9" but real release notes rarely go past 4.
HEADING_STYLES = {
    1: "Heading 1",
    2: "Heading 2",
    3: "Heading 3",
    4: "Heading 4",
    5: "Heading 5",
    6: "Heading 6",
}
DEFAULT_BODY_STYLE = "Body Text"
DEFAULT_CAPTION_STYLE = "Caption"
DEFAULT_BULLET_STYLE = "Bullet 1"
DEFAULT_NUMBERED_STYLE = "Step 1"
TABLE_HEADING_STYLE = "Table Heading"
TABLE_TEXT_STYLE = "Table Text"

# Element types we deliberately skip when walking the top-level kids tree.
# Headers/footers are page-repeating chrome (filename, page number, copyright)
# that would clutter the body if we transcribed every page's copy.
SKIP_TYPES = {"header", "footer"}

# Cover-page text-run substitutions. The template's cover page lives inside
# drawing/text-box elements that python-docx does not surface via doc.paragraphs,
# so we walk every <w:t> element in the document body and replace by exact match.
# Verified against the actual template at template-discovery time — the literal
# strings on the left are the exact contents of the cover-page text runs.

# Phrases that mark a text-box paragraph as instructional clutter we want to
# wipe entirely. Matched as substrings against the joined paragraph text so
# they catch fragmented runs (e.g. "If this is your first ti" + "m" + "e ...").
COVER_INSTRUCTION_PHRASES = (
    "If this is your first ti",
    "Click on this text box and delete it",
    "Click this text box and delete it",
    "Update Customer Name",
)

# Cap embedded image width so wide screenshots don't blow past the page margins.
MAX_IMAGE_WIDTH = Inches(5.5)

# Extract a version like 8.1.12.0 from the PDF filename.
VERSION_RE = re.compile(r"(\d+(?:\.\d+){1,3})")

# Strip leading bullet glyphs from list-item text. The extractor includes the
# literal bullet character in `content` (e.g. "• (AM1393) ..."), but the Word
# Bullet 1 style draws its own bullet — keeping the glyph would double up.
LEADING_BULLET_RE = re.compile(r"^\s*[•·●▪◦‣⁃-]\s*")

# Detect TOC-leader rows. Release-notes PDFs include a tiny on-page TOC on
# page 1 that the extractor parses as headings/list items, e.g.
#     "AM3394: [FFT] – Security improvements............................2"
# These duplicate the real content on later pages and add noise. Match three
# or more dots followed by a trailing page number.
TOC_LEADER_RE = re.compile(r"\.{3,}\s*\d+\s*$")

# Detect TOC-style headings without dotted leaders. The cover-page TOC also
# emits short headings like "Release Features 2" / "Defect Fixes 4" where the
# trailing number is the page reference. Match: short string ending in space +
# 1-3 digit number, with no other digits in the text. The "no other digits"
# guard avoids stripping legitimate headings like "Section 1.2.3 Overview".
TOC_TRAILING_NUM_RE = re.compile(r"^([^\d]{1,60}?)\s+(\d{1,3})$")

# Pattern-based release-notes promotion. The fast extractor judges heading
# level by font size and gets it wrong on every release-notes PDF we've seen
# (Release Features as L4, Defect Fixes as L3, AM items as bullets/paragraphs).
# We override its decisions for known release-notes structures.
#
# The example output DOCX (`docs_example/output_docx_example/...`) shows the
# canonical hierarchy CAE expects:
#   Heading 1 = section name (Release Features / Defect Fixes / Not Tested ...)
#   Heading 2 = each AM item (AM####: <title>)
# The Word TOC field rebuilds itself from these styles.

# Section names that should always become Heading 1.
RELEASE_SECTION_NAMES = frozenset({
    "release features",
    "defect fixes",
    "not tested",
    "miscellaneous",
    "introduction",
    "remark",
    "remarks",
    "known issues",
    "known limitations",
    "compatibility",
    "installation",
    "release notes",
})

# Match an AM-item line: leading bullet glyph + "AM" + digits + separator.
# Allows trailing colon, hyphen, space, period, dash, or bracket variants.
AM_ITEM_RE = re.compile(r"^\s*[•·●▪◦‣⁃-]?\s*(AM\d{2,5})\s*[:\-–—]")

# Sub-headings inside an AM item ("Bug Description:", "After correction:",
# "Steps to reproduce:", "Workaround:", etc.). These are short labels ending
# in a colon and should sit one level below an AM heading (Heading 3).
AM_SUBHEADING_NAMES = frozenset({
    "bug description",
    "steps to reproduce",
    "expected result",
    "actual result",
    "after correction",
    "workaround",
    "root cause",
    "impact",
    "fix description",
    "description",
    "remark",
    "remarks",
    "notes",
})

# Page chrome that hybrid mode (Docling) emits as plain paragraphs/headings
# instead of header/footer elements. Match by exact substring against the
# trimmed text. These repeat on every page and add no value to the body.
PAGE_CHROME_PHRASES = (
    "Jetsched Communications Release Note",
    "AIR TRANSPORT INFORMATION SYSTEMS",
    "SAS CYBERJET",
    "RCS BORDEAUX",
    "Office :",
    "Page ",  # "Page X sur 4"
)

# Bare metadata fragments the page-header band emits: a version triple and a
# date in MM/YYYY format. These appear without context as their own paragraphs.
PAGE_CHROME_RE = re.compile(
    r"^\s*("
    r"\d+\.\d+\.\d+(?:\.\d+)?"          # version triple, e.g. 8.1.12.0
    r"|\d{1,2}/\d{4}"                    # MM/YYYY date
    r")\s*$"
)

logger = logging.getLogger("scripts.docx_conversion")


# ─────────────────────────── logging ───────────────────────────


def setup_logging(verbose: bool) -> None:
    logging.basicConfig(
        level=logging.INFO,
        format="[%(asctime)s] [%(levelname)s] %(message)s",
        datefmt="%H:%M:%S",
    )
    # --verbose only enables DEBUG for our own loggers, not third-party libs
    if verbose:
        for name in ("scripts.docx_conversion", "claude.client", "claude.extractor",
                      "pdf.image_extractor"):
            logging.getLogger(name).setLevel(logging.DEBUG)


# ─────────────────────────── extraction ───────────────────────────


def _hybrid_server_up() -> bool:
    """Quick TCP probe — is the opendataloader-pdf-hybrid server listening?"""
    with socket.socket(socket.AF_INET, socket.SOCK_STREAM) as s:
        s.settimeout(0.5)
        try:
            s.connect((HYBRID_HOST, HYBRID_PORT))
            return True
        except OSError:
            return False


def resolve_mode(requested: str) -> str:
    """Pick the extraction mode, falling back to fast if hybrid prerequisites
    aren't met. Never silently switches without an explicit log line."""
    if requested == "fast":
        return "fast"
    if requested == "hybrid":
        if not _hybrid_server_up():
            logger.warning(
                "convert.mode.fallback requested=hybrid reason=hybrid_server_down "
                "host=%s port=%d → falling back to fast. "
                "Start the server with: opendataloader-pdf-hybrid --port %d",
                HYBRID_HOST, HYBRID_PORT, HYBRID_PORT,
            )
            return "fast"
        return "hybrid"
    if requested == "claude":
        return "claude"
    raise ValueError(f"Unknown extraction mode: {requested}")


def extract_pdf(pdf_path: Path, mode: str, use_cache: bool) -> Path:
    """Run opendataloader-pdf on `pdf_path` and return the path to the JSON
    output. Caches results under docs_example/conversion_prototype/.cache/<mode>/
    so re-runs are instant.
    """
    cache_dir = CACHE_ROOT / mode
    cache_dir.mkdir(parents=True, exist_ok=True)
    json_path = cache_dir / f"{pdf_path.stem}.json"

    if use_cache and json_path.exists() and json_path.stat().st_size > 0:
        logger.info(
            "extract.cache.hit pdf=%s mode=%s json=%s",
            pdf_path.name, mode, json_path.relative_to(PROJECT_ROOT),
        )
        return json_path

    logger.info("extract.start pdf=%s mode=%s", pdf_path.name, mode)
    kwargs: dict[str, Any] = {
        "input_path": [str(pdf_path)],
        "output_dir": str(cache_dir),
        "format": "markdown,json",
    }
    if mode == "hybrid":
        # The "docling-fast" hybrid backend talks to the locally-running
        # opendataloader-pdf-hybrid server (IBM Docling under the hood).
        kwargs["hybrid"] = "docling-fast"

    opendataloader_pdf.convert(**kwargs)

    if not json_path.exists():
        raise RuntimeError(
            f"opendataloader-pdf produced no JSON for {pdf_path.name} in {cache_dir}"
        )
    logger.info(
        "extract.success pdf=%s mode=%s json=%s size=%d",
        pdf_path.name, mode, json_path.relative_to(PROJECT_ROOT), json_path.stat().st_size,
    )
    return json_path


# ─────────────────────────── claude mode ───────────────────────────


def _claude_deps():
    """Lazy-import Claude extractor + image extractor from backend/.

    Only called for ``--mode claude`` so the fast/hybrid modes don't need
    the backend package on sys.path.
    """
    sys.path.insert(0, str(PROJECT_ROOT / "backend"))
    from app.integrations.pdf.image_extractor import extract_images  # noqa: E402
    from app.integrations.claude.extractor import extract_release_note  # noqa: E402
    from app.integrations.claude.client import ClaudeClient  # noqa: E402
    return extract_images, extract_release_note, ClaudeClient


def _claude_cache_path(pdf_path: Path) -> Path:
    """Cache path keyed by PDF content hash."""
    pdf_hash = hashlib.sha256(pdf_path.read_bytes()).hexdigest()
    return CLAUDE_CACHE_ROOT / f"{pdf_hash}.json"


def _load_claude_cache(pdf_path: Path):
    """Load a cached ReleaseNoteRecord if it exists. Returns None on miss."""
    cache_file = _claude_cache_path(pdf_path)
    if not cache_file.exists():
        return None
    try:
        sys.path.insert(0, str(PROJECT_ROOT / "backend"))
        from app.state.release_notes_models import ReleaseNoteRecord  # noqa: E402
        record = ReleaseNoteRecord.model_validate_json(cache_file.read_text())
        logger.info("claude.cache.hit pdf=%s cache=%s", pdf_path.name, cache_file.name[:12])
        return record
    except Exception as exc:
        logger.warning("claude.cache.invalid pdf=%s error=%s", pdf_path.name, exc)
        return None


def _save_claude_cache(pdf_path: Path, record) -> None:
    """Persist a ReleaseNoteRecord to the Claude cache."""
    CLAUDE_CACHE_ROOT.mkdir(parents=True, exist_ok=True)
    cache_file = _claude_cache_path(pdf_path)
    cache_file.write_text(record.model_dump_json(indent=2))
    logger.info("claude.cache.saved pdf=%s cache=%s", pdf_path.name, cache_file.name[:12])


def format_item_heading(item) -> str:
    """Build the Heading 2 text for a release-note item.

    Format: ``AM1393 [HAL] - Adding characters replacement feature``
    """
    heading = item.am_card
    if item.customers:
        heading += f" [{', '.join(item.customers)}]"
    heading += f" - {item.title}"
    return heading


def render_record(doc: Document, record, images_dir: Path) -> dict[str, int]:
    """Walk a ReleaseNoteRecord and emit content into the DOCX.

    Replaces ``normalize_heading_levels`` + ``collect_image_bboxes`` +
    ``walk_document`` for claude mode. Returns a counts dict for logging.
    """
    counts: dict[str, int] = {"sections": 0, "items": 0, "blocks": 0}

    # Group items by section, preserving first-seen order
    sections: dict[str, list] = {}
    for item in record.items:
        sections.setdefault(item.section, []).append(item)

    for section_name, items in sections.items():
        add_styled_paragraph(doc, section_name, "Heading 1")
        counts["sections"] += 1

        for item in items:
            add_styled_paragraph(doc, format_item_heading(item), "Heading 2")
            counts["items"] += 1

            if item.summary:
                add_styled_paragraph(doc, item.summary, DEFAULT_BODY_STYLE)

            for block in item.body:
                counts["blocks"] += 1
                btype = block.type

                if btype == "paragraph":
                    add_styled_paragraph(doc, block.text, DEFAULT_BODY_STYLE)

                elif btype == "heading":
                    add_bold_body_paragraph(doc, block.text)

                elif btype == "image":
                    img_path = images_dir / f"{block.image_id}.png"
                    if img_path.exists():
                        try:
                            doc.add_picture(str(img_path), width=MAX_IMAGE_WIDTH)
                        except Exception as exc:
                            logger.warning("image.embed.failed id=%s error=%s", block.image_id, exc)
                    else:
                        logger.warning("image.missing id=%s path=%s", block.image_id, img_path)

                elif btype == "list":
                    style = DEFAULT_NUMBERED_STYLE if block.ordered else DEFAULT_BULLET_STYLE
                    for li_text in block.items:
                        add_styled_paragraph(doc, li_text, style)

                elif btype == "table":
                    if block.headers or block.rows:
                        n_cols = len(block.headers) if block.headers else (
                            max((len(r) for r in block.rows), default=0)
                        )
                        total_rows = (1 if block.headers else 0) + len(block.rows)
                        if n_cols > 0 and total_rows > 0:
                            table = doc.add_table(rows=total_rows, cols=n_cols)
                            row_idx = 0
                            if block.headers:
                                for c, h in enumerate(block.headers):
                                    cell = table.rows[0].cells[c]
                                    cell.paragraphs[0].text = h
                                    if style_safe(doc, TABLE_HEADING_STYLE):
                                        cell.paragraphs[0].style = doc.styles[TABLE_HEADING_STYLE]
                                row_idx = 1
                            for r in block.rows:
                                for c, val in enumerate(r):
                                    if c < n_cols:
                                        cell = table.rows[row_idx].cells[c]
                                        cell.paragraphs[0].text = val
                                        if style_safe(doc, TABLE_TEXT_STYLE):
                                            cell.paragraphs[0].style = doc.styles[TABLE_TEXT_STYLE]
                                row_idx += 1

                elif btype == "code":
                    add_styled_paragraph(doc, block.text, DEFAULT_BODY_STYLE)

    return counts


# ─────────────────────────── template prep ───────────────────────────


def derive_version(pdf_path: Path) -> str:
    m = VERSION_RE.search(pdf_path.stem)
    return m.group(1) if m else "Unknown"


def cover_replacements(product: str, version: str) -> dict[str, str]:
    """Build the literal text-run substitutions for the cover page."""
    return {
        "External Business Document": "Release Notes",
        "Customer": "ACARS",  # category line above the product name
        "Sample Product Name": product,
        "Version #": f"Version {version}",
        "Date": datetime.now().strftime("%B %Y"),
    }


def patch_cover_page(doc: Document, replacements: dict[str, str]) -> int:
    """Walk every <w:t> element in the document body (including text boxes
    that python-docx hides from doc.paragraphs) and apply exact-match
    substitutions. Returns the number of runs replaced.
    """
    replaced = 0
    for t in doc.element.body.iter(qn("w:t")):
        if t.text in replacements:
            old = t.text
            t.text = replacements[old]
            replaced += 1
            logger.debug("cover.replace old=%r new=%r", old, t.text)
    return replaced


def clean_cover_textboxes(doc: Document) -> int:
    """Wipe instructional text inside cover-page text boxes (`<w:txbxContent>`).

    The Flightscape cover page has small "If this is your first time…" /
    "Click this text box and delete it" hint paragraphs that python-docx
    can't reach via `doc.paragraphs`. We walk every text-box paragraph,
    join its <w:t> runs, and if the joined text matches any instruction
    phrase, blank out every run in that paragraph. Returns paragraphs cleaned.
    """
    cleaned = 0
    for txbx in doc.element.body.iter(qn("w:txbxContent")):
        for p in txbx.iter(qn("w:p")):
            runs = list(p.iter(qn("w:t")))
            joined = "".join(r.text or "" for r in runs)
            if not joined.strip():
                continue
            if any(phrase in joined for phrase in COVER_INSTRUCTION_PHRASES):
                for r in runs:
                    r.text = ""
                cleaned += 1
                logger.debug("cover.clean joined=%r", joined[:80])
    return cleaned


def mark_toc_dirty(doc: Document) -> bool:
    """Find the TOC field in the body and mark it dirty so Word rebuilds it
    from the document's actual headings the next time the file is opened.

    OOXML stores a Table of Contents as a complex field bracketed by
    `<w:fldChar w:fldCharType="begin"/>` ... `<w:fldChar w:fldCharType="end"/>`
    with the field instructions in `<w:instrText>` (e.g. ``TOC \\o "2-5" ...``).
    The cached visible content between `separate` and `end` is what shows
    the stale "Insert your Heading" / "Delete This Chapter" entries from
    the original template. Adding `w:dirty="1"` to the begin fldChar tells
    Word "this field's cache is out of date" — when the user opens the
    document Word offers to update fields (or auto-updates if configured).

    We can't easily replace the cached entries programmatically without
    breaking the bookmarks/PAGEREFs the field uses internally.
    """
    body = doc.element.body
    fldChar_tag = qn("w:fldChar")
    instrText_tag = qn("w:instrText")
    fldCharType_attr = qn("w:fldCharType")
    dirty_attr = qn("w:dirty")

    # Walk the body and find the begin fldChar that opens a TOC field. The
    # TOC instrText follows the begin marker (possibly split across multiple
    # instrText elements: 'TOC \\o "2-5" \\' + 'n 1-1 \\' + ...).
    elements = list(body.iter())
    for i, elem in enumerate(elements):
        if elem.tag != fldChar_tag:
            continue
        if elem.get(fldCharType_attr) != "begin":
            continue
        # Look ahead for the next instrText to see what kind of field this is.
        for j in range(i + 1, min(i + 6, len(elements))):
            nxt = elements[j]
            if nxt.tag == instrText_tag and (nxt.text or "").lstrip().startswith("TOC"):
                elem.set(dirty_attr, "1")
                logger.debug("toc.dirty.marked instrText=%r", (nxt.text or "")[:40])
                return True
            if nxt.tag == fldChar_tag:
                break  # different field
    return False


def _has_sectpr(p) -> bool:
    """True if this paragraph carries a sectPr (section break), which means
    deleting it would collapse the page layout. We never touch these."""
    pPr = p._p.find(qn("w:pPr"))
    return pPr is not None and pPr.find(qn("w:sectPr")) is not None


# Index of the first body paragraph in the unmodified Flightscape template.
# Verified by inspection: paragraphs 0-37 = cover page (with anchored drawings,
# pictures, and text boxes) + TOC + section breaks; paragraph 38 onwards is
# the instructional example chapters we strip. The cover-page artwork lives
# inside drawings ANCHORED to paragraphs 0 and 5 — deleting those paragraphs
# destroys the artwork, so we never touch any paragraph with index < 38.
TEMPLATE_BODY_START_INDEX = 38


def strip_template_body(doc: Document) -> tuple[int, int]:
    """Remove the template's instructional example chapters (paragraphs 38+
    and any tables that live after them) while preserving the cover page,
    the TOC, and all section breaks.

    Strategy: do a single sequential walk over the body's child elements
    (paragraphs and tables interleaved). Count paragraphs as we go; once we
    pass `TEMPLATE_BODY_START_INDEX`, delete every paragraph (unless it's a
    section-break carrier) and every table we encounter. This preserves
    cover-page drawings/pictures/text boxes that are anchored inside the
    early paragraphs.
    """
    body = doc.element.body
    p_tag = qn("w:p")
    tbl_tag = qn("w:tbl")
    sectPr_tag = qn("w:sectPr")

    # Top-level w:sectPr (the final section properties) — never delete this.
    # It's a direct child of body, not nested inside a paragraph.

    para_index = 0
    removed_paras = 0
    removed_tables = 0

    # Snapshot children first; we'll mutate `body` while iterating the snapshot.
    for child in list(body):
        tag = child.tag
        if tag == p_tag:
            current_index = para_index
            para_index += 1
            if current_index < TEMPLATE_BODY_START_INDEX:
                continue  # cover page or TOC — keep as-is

            # Past the body cutoff. Keep section-break carriers; delete the rest.
            pPr = child.find(qn("w:pPr"))
            if pPr is not None and pPr.find(qn("w:sectPr")) is not None:
                continue
            body.remove(child)
            removed_paras += 1
        elif tag == tbl_tag:
            # Tables before the cutoff are part of the cover/TOC layout — keep.
            # Tables after the cutoff are instructional examples — delete.
            if para_index < TEMPLATE_BODY_START_INDEX:
                continue
            body.remove(child)
            removed_tables += 1
        elif tag == sectPr_tag:
            # Top-level sectPr (final section properties) — never delete.
            continue
        # Anything else (bookmarks, etc.) — leave untouched.

    logger.debug(
        "strip.body removed_paras=%d removed_tables=%d total_paras_seen=%d",
        removed_paras, removed_tables, para_index,
    )
    return removed_paras, removed_tables


# ─────────────────────────── content emission ───────────────────────────


def style_safe(doc: Document, name: str) -> str | None:
    """Return `name` if it exists in the doc's style table, else None.
    python-docx raises KeyError when applying an unknown style — we want a
    soft fallback to default formatting instead."""
    try:
        doc.styles[name]
        return name
    except KeyError:
        return None


def add_styled_paragraph(doc: Document, text: str, style: str | None) -> None:
    if not text or not text.strip():
        return
    p = doc.add_paragraph(text)
    if style and style_safe(doc, style):
        p.style = doc.styles[style]


def add_bold_body_paragraph(doc: Document, text: str) -> None:
    """Body Text paragraph with the entire run bolded. Used for AM
    sub-headings (Bug Description:, After correction:, …) so they're
    visually distinct without polluting the TOC."""
    if not text or not text.strip():
        return
    p = doc.add_paragraph()
    if style_safe(doc, DEFAULT_BODY_STYLE):
        p.style = doc.styles[DEFAULT_BODY_STYLE]
    run = p.add_run(text)
    run.bold = True


def _is_junk_text(text: str) -> bool:
    """True if `text` is meaningless body content we should drop.

    The extractor occasionally surfaces page numbers and similar 1-2 char
    fragments as standalone paragraphs ("1", "2", "i", "iv", etc.) and
    duplicates real content as TOC-leader rows on the cover page. Hybrid
    mode also leaks page header/footer chrome (CAE address block, version
    triples, "Page X sur N") as flat paragraphs."""
    s = (text or "").strip()
    if not s:
        return True
    if len(s) <= 3 and (s.isdigit() or s.lower() in {"i", "ii", "iii", "iv", "v", "vi"}):
        return True
    if TOC_LEADER_RE.search(s):
        return True
    if PAGE_CHROME_RE.match(s):
        return True
    if any(phrase in s for phrase in PAGE_CHROME_PHRASES):
        return True
    return False


def _is_toc_heading(text: str) -> bool:
    """True if a heading looks like a cover-page TOC entry: short text with
    a trailing standalone page number, e.g. "Release Features 2"."""
    return bool(TOC_TRAILING_NUM_RE.match((text or "").strip()))


def classify_release_note_line(text: str) -> str | None:
    """Pattern-based classification for release-notes content. Returns the
    Word style name a line should use, or None if no pattern matches.

    Two rules, in priority order:
      1. Known section name (Release Features / Defect Fixes / Not Tested
         / ...) → Heading 1.
      2. AM-code line (AM3394: …) → Heading 2.

    AM sub-headings (Bug Description:, After correction:, etc.) are NOT
    promoted to a heading style — they should appear inline as bold
    labels in the body, not as TOC entries. They're handled separately
    by `is_am_subheading()` so the emitter can apply bold formatting
    without contributing to the Word TOC field.

    These overrides apply regardless of how the extractor classified the
    element — paragraphs, list items, headings, captions all get
    pattern-promoted when they match.
    """
    s = (text or "").strip().rstrip(" :.")
    if not s:
        return None
    lower = s.lower()
    if lower in RELEASE_SECTION_NAMES:
        return "Heading 1"
    if AM_ITEM_RE.match(text or ""):
        return "Heading 2"
    return None


def is_am_subheading(text: str) -> bool:
    """True if `text` is an inline AM sub-heading label like "Bug Description:".

    These should render as bold body text — visually distinct, but NOT
    captured by the Word TOC field (only Heading 1/2/3 styles feed it)."""
    s = (text or "").strip().rstrip(" :.")
    return s.lower() in AM_SUBHEADING_NAMES


def collect_image_bboxes(payload: dict[str, Any]) -> dict[int, list[list[float]]]:
    """Walk the JSON tree and group every image's bounding box by page number.

    Hybrid mode (Docling) runs OCR over embedded screenshots and emits the
    OCR'd UI labels as paragraphs that sit *inside* the image's bounding box
    on the same page. Those paragraphs are noise — they're already represented
    visually by the embedded image. The walker uses this map to drop any text
    element whose center falls inside an image's bbox on the same page.

    Logo images (top-of-page Cyberjet icon) are excluded from this map: they
    overlap the page-header chrome filter handles separately, and we don't
    want to accidentally drop real text that sits just below the logo.
    """
    by_page: dict[int, list[list[float]]] = {}

    def walk(node: Any) -> None:
        if isinstance(node, dict):
            if node.get("type") == "image" and not is_page_header_logo(node):
                pg = node.get("page number")
                bb = node.get("bounding box")
                if isinstance(pg, int) and isinstance(bb, list) and len(bb) == 4:
                    by_page.setdefault(pg, []).append([float(x) for x in bb])
            for v in node.values():
                walk(v)
        elif isinstance(node, list):
            for item in node:
                walk(item)

    walk(payload)
    return by_page


def is_inside_image(node: dict[str, Any], image_bboxes: dict[int, list[list[float]]]) -> bool:
    """True if `node`'s bbox center sits inside any image bbox on the same page."""
    pg = node.get("page number")
    bb = node.get("bounding box") or []
    if not isinstance(pg, int) or len(bb) != 4:
        return False
    cx = (bb[0] + bb[2]) / 2
    cy = (bb[1] + bb[3]) / 2
    for ibox in image_bboxes.get(pg, []):
        if ibox[0] <= cx <= ibox[2] and ibox[1] <= cy <= ibox[3]:
            return True
    return False


def normalize_heading_levels(payload: dict[str, Any]) -> dict[int, int]:
    """Walk the JSON tree once and build a mapping from extractor heading
    levels (which are font-size derived and inconsistent — e.g. {2,3,4,5,7})
    to a contiguous Word level range (1, 2, 3, 4, 5, …).

    The smallest extractor level becomes Heading 1, the next smallest
    becomes Heading 2, and so on. This keeps the relative hierarchy intact
    while producing a clean Heading 1/2/3 outline in Word."""
    seen: set[int] = set()

    def walk(node: Any) -> None:
        if isinstance(node, dict):
            if node.get("type") == "heading":
                text = node.get("content", "") or ""
                # Skip headings we'll drop anyway so they don't pollute the
                # level mapping.
                if not _is_junk_text(text) and not _is_toc_heading(text):
                    lvl = node.get("heading level")
                    if isinstance(lvl, int):
                        seen.add(lvl)
            for v in node.values():
                walk(v)
        elif isinstance(node, list):
            for item in node:
                walk(item)

    walk(payload)
    sorted_levels = sorted(seen)
    mapping = {orig: i + 1 for i, orig in enumerate(sorted_levels)}
    logger.debug("heading.normalize observed=%s mapping=%s", sorted_levels, mapping)
    return mapping


def emit_paragraph(doc: Document, node: dict[str, Any]) -> None:
    text = node.get("content", "")
    if _is_junk_text(text):
        return
    if is_am_subheading(text):
        add_bold_body_paragraph(doc, text)
        return
    promoted = classify_release_note_line(text)
    add_styled_paragraph(doc, text, promoted or DEFAULT_BODY_STYLE)


def emit_heading(
    doc: Document,
    node: dict[str, Any],
    level_map: dict[int, int],
) -> None:
    text = node.get("content", "") or ""
    if _is_junk_text(text) or _is_toc_heading(text):
        return
    if is_am_subheading(text):
        add_bold_body_paragraph(doc, text)
        return
    promoted = classify_release_note_line(text)
    if promoted:
        add_styled_paragraph(doc, text, promoted)
        return
    raw_level = int(node.get("heading level") or 1)
    level = level_map.get(raw_level, raw_level)
    style = HEADING_STYLES.get(level, "Heading 6")
    add_styled_paragraph(doc, text, style)


def emit_caption(doc: Document, node: dict[str, Any]) -> None:
    text = node.get("content", "") or ""
    promoted = classify_release_note_line(text)
    add_styled_paragraph(doc, text, promoted or DEFAULT_CAPTION_STYLE)


def emit_list(
    doc: Document,
    node: dict[str, Any],
    level_map: dict[int, int],
    image_bboxes: dict[int, list[list[float]]],
    image_root: Path | None,
) -> None:
    numbering_style = (node.get("numbering style") or "unordered").lower()
    bullet_style = (
        DEFAULT_NUMBERED_STYLE if numbering_style.startswith("ordered") else DEFAULT_BULLET_STYLE
    )
    for item in node.get("list items") or []:
        raw = item.get("content") or ""
        # Promote AM-item lines BEFORE stripping the bullet glyph — the
        # classifier matches against the original text including the bullet.
        promoted = classify_release_note_line(raw)
        text = LEADING_BULLET_RE.sub("", raw)
        if _is_junk_text(text):
            continue
        add_styled_paragraph(doc, text, promoted or bullet_style)
        # Some list items also carry kids (nested content) — flatten them.
        for child in item.get("kids") or []:
            walk_node(doc, child, level_map, image_bboxes, image_root)


def _flatten_cell_text(cell: dict[str, Any]) -> str:
    """A table cell's `kids` are paragraphs/lists/etc. The Flightscape
    Table Text style is paragraph-level, so we collapse a cell's content to
    a single string and let the cell's style handle wrapping."""
    parts: list[str] = []
    for kid in cell.get("kids") or []:
        if isinstance(kid, dict):
            content = kid.get("content")
            if content:
                parts.append(content)
            # Lists inside cells: pull each item's content too.
            for li in kid.get("list items") or []:
                lc = li.get("content")
                if lc:
                    parts.append(LEADING_BULLET_RE.sub("", lc))
    return "\n".join(parts).strip()


def emit_table(doc: Document, node: dict[str, Any]) -> None:
    rows = node.get("rows") or []
    n_cols = int(node.get("number of columns") or 0) or max(
        (len(r.get("cells") or []) for r in rows), default=0
    )
    if not rows or n_cols == 0:
        return

    table = doc.add_table(rows=len(rows), cols=n_cols)
    # The Flightscape "Table Grid"-equivalent style isn't strictly necessary;
    # we set per-cell paragraph styles instead so the cells inherit fonts.
    for r_idx, row in enumerate(rows):
        cells = row.get("cells") or []
        for c_idx in range(n_cols):
            target_cell = table.rows[r_idx].cells[c_idx]
            text = ""
            if c_idx < len(cells):
                text = _flatten_cell_text(cells[c_idx])
            # Reuse the auto-created first paragraph; clear it then style it.
            cell_para = target_cell.paragraphs[0]
            cell_para.text = text
            style_name = TABLE_HEADING_STYLE if r_idx == 0 else TABLE_TEXT_STYLE
            if style_safe(doc, style_name):
                cell_para.style = doc.styles[style_name]


def is_page_header_logo(node: dict[str, Any]) -> bool:
    """True if this image looks like the small Cyberjet logo embedded in the
    top-left corner of every page header.

    The logo is consistently ~72×72pt at roughly (x=67, y_top=761) — i.e. it
    lives in the top page-header band and is much smaller than any real
    screenshot in the body. We filter it out so the generated DOCX doesn't
    repeat the corporate icon between every section.

    Heuristic: image whose bbox top sits above y=680 (top inch of an A4 page)
    AND is under 200×200pt. Real release-notes screenshots are always wider
    than that and live further down the page.
    """
    bb = node.get("bounding box") or []
    if len(bb) != 4:
        return False
    x0, y0, x1, y1 = bb
    width = x1 - x0
    height = y1 - y0
    return y1 > 680 and width < 200 and height < 200


def emit_image(doc: Document, node: dict[str, Any], image_root: Path) -> None:
    if is_page_header_logo(node):
        logger.debug("image.skip.logo bbox=%s", node.get("bounding box"))
        return
    rel = node.get("source")
    if not rel:
        return
    img_path = (image_root / rel).resolve()
    if not img_path.exists():
        logger.debug("image.skip.missing path=%s", img_path)
        return
    try:
        doc.add_picture(str(img_path), width=MAX_IMAGE_WIDTH)
    except Exception as exc:
        # Some PDFs embed unsupported image formats (CMYK JPEGs, for instance).
        # Don't fail the whole conversion — log and continue.
        logger.warning("image.embed.failed path=%s error=%s", img_path.name, exc)


def walk_node(
    doc: Document,
    node: Any,
    level_map: dict[int, int],
    image_bboxes: dict[int, list[list[float]]],
    image_root: Path | None = None,
) -> None:
    """Top-level dispatch for one element from the opendataloader JSON tree."""
    if not isinstance(node, dict):
        return
    t = node.get("type")
    if t in SKIP_TYPES:
        return
    # Drop any text element that lives inside an image's bbox — that's OCR
    # noise from hybrid mode running over an embedded screenshot.
    if t in {"heading", "paragraph", "caption"} and is_inside_image(node, image_bboxes):
        return
    if t == "heading":
        emit_heading(doc, node, level_map)
    elif t == "paragraph":
        emit_paragraph(doc, node)
    elif t == "caption":
        emit_caption(doc, node)
    elif t == "list":
        emit_list(doc, node, level_map, image_bboxes, image_root)
    elif t == "table":
        emit_table(doc, node)
    elif t == "image":
        if image_root is not None:
            emit_image(doc, node, image_root)
    else:
        # Unknown / container type — descend into kids if any.
        for child in node.get("kids") or []:
            walk_node(doc, child, level_map, image_bboxes, image_root)


def walk_document(
    doc: Document,
    payload: dict[str, Any],
    image_root: Path,
    level_map: dict[int, int],
    image_bboxes: dict[int, list[list[float]]],
) -> dict[str, int]:
    """Walk the top-level kids list and emit each element. Returns a count
    breakdown for the run summary."""
    counts: dict[str, int] = {}
    for child in payload.get("kids") or []:
        if isinstance(child, dict):
            t = child.get("type") or "unknown"
            counts[t] = counts.get(t, 0) + 1
        walk_node(doc, child, level_map, image_bboxes, image_root)
    return counts


# ─────────────────────────── orchestration ───────────────────────────


def convert(
    pdf_path: Path,
    template_path: Path,
    output_path: Path,
    *,
    mode: str,
    product: str,
    use_cache: bool,
    prebuilt_json: Path | None = None,
) -> None:
    if not pdf_path.exists():
        raise FileNotFoundError(f"PDF not found: {pdf_path}")
    if not template_path.exists():
        raise FileNotFoundError(f"Template not found: {template_path}")

    version = derive_version(pdf_path)
    logger.info(
        "convert.start pdf=%s version=%s product=%r mode=%s",
        pdf_path.name, version, product, mode,
    )

    # ── Extraction ────────────────────────────────────────────────
    if mode == "claude":
        extract_images_fn, extract_release_note_fn, ClaudeClient = _claude_deps()
        record = None
        if use_cache:
            record = _load_claude_cache(pdf_path)
        if record is None:
            manifest = extract_images_fn(pdf_path)
            api_key = os.environ.get("ANTHROPIC_API_KEY", "")
            client = ClaudeClient(api_key)
            record = extract_release_note_fn(
                pdf_path, manifest, version=version, claude_client=client,
            )
            _save_claude_cache(pdf_path, record)
        images_dir = pdf_path.parent / "images"
    else:
        if prebuilt_json is not None:
            if not prebuilt_json.exists():
                raise FileNotFoundError(f"--json path not found: {prebuilt_json}")
            json_path = prebuilt_json.resolve()
            try:
                json_rel = json_path.relative_to(PROJECT_ROOT)
            except ValueError:
                json_rel = json_path
            logger.info("extract.skipped reason=prebuilt_json json=%s", json_rel)
        else:
            json_path = extract_pdf(pdf_path, mode, use_cache)
        payload = json.loads(json_path.read_text(encoding="utf-8"))
        image_root = json_path.parent

    # ── Shared template prep ─────────────────────────────────────
    doc = Document(str(template_path))
    cover_count = patch_cover_page(doc, cover_replacements(product, version))
    logger.info("cover.patched runs=%d", cover_count)

    cleaned = clean_cover_textboxes(doc)
    logger.info("cover.textboxes.cleaned paragraphs=%d", cleaned)

    stripped_paras, stripped_tables = strip_template_body(doc)
    logger.info(
        "template.body.stripped paragraphs=%d tables=%d",
        stripped_paras, stripped_tables,
    )

    if mark_toc_dirty(doc):
        logger.info("toc.dirty.marked status=ok")
    else:
        logger.warning("toc.dirty.marked status=not_found")

    # ── Rendering ────────────────────────────────────────────────
    if mode == "claude":
        counts = render_record(doc, record, images_dir)
        logger.info("convert.body.emitted breakdown=%s", counts)
    else:
        level_map = normalize_heading_levels(payload)
        logger.info("heading.normalize map=%s", level_map)

        image_bboxes = collect_image_bboxes(payload)
        total_images = sum(len(v) for v in image_bboxes.values())
        logger.info("image.bboxes.collected pages=%d total=%d", len(image_bboxes), total_images)

        counts = walk_document(doc, payload, image_root, level_map, image_bboxes)
        logger.info(
            "convert.body.emitted top_level_elements=%d breakdown=%s",
            sum(counts.values()), counts,
        )

    # ── Save ─────────────────────────────────────────────────────
    output_path.parent.mkdir(parents=True, exist_ok=True)
    doc.save(str(output_path))
    saved = output_path.resolve()
    try:
        rel = saved.relative_to(PROJECT_ROOT)
    except ValueError:
        rel = saved
    logger.info(
        "convert.success pdf=%s output=%s size=%d",
        pdf_path.name, rel, saved.stat().st_size,
    )

    # Claude mode: save the full record JSON next to the DOCX for inspection.
    if mode == "claude":
        record_path = output_path.with_suffix(".json")
        record_path.write_text(record.model_dump_json(indent=2))
        try:
            rec_rel = record_path.resolve().relative_to(PROJECT_ROOT)
        except ValueError:
            rec_rel = record_path.resolve()
        logger.info("convert.record.saved path=%s", rec_rel)

    # Drop the extractor's Markdown next to the DOCX (fast/hybrid only —
    # claude mode doesn't produce an opendataloader markdown).
    if mode != "claude":
        json_path_local = json_path  # noqa: F841 — defined in the else branch above
        md_source = json_path.with_suffix(".md")
        if md_source.exists():
            md_target = output_path.with_suffix(".md")
            shutil.copy2(md_source, md_target)
            try:
                md_rel = md_target.resolve().relative_to(PROJECT_ROOT)
            except ValueError:
                md_rel = md_target.resolve()
            logger.info("convert.md.copied source=%s target=%s", md_source.name, md_rel)
        else:
            logger.warning("convert.md.missing source=%s", md_source)


# ─────────────────────────── CLI ───────────────────────────


def parse_args() -> argparse.Namespace:
    p = argparse.ArgumentParser(
        description="PDF → DOCX conversion prototype (Unit 4 Block B gate)",
    )
    p.add_argument("--pdf", required=True, type=Path,
                   help="Path to a release-notes PDF (e.g. patches/.../X.Y.Z - Release Notes.pdf)")
    p.add_argument("--output", required=True, type=Path,
                   help="Where to write the generated DOCX")
    p.add_argument("--template", default=DEFAULT_TEMPLATE, type=Path,
                   help=f"Flightscape template path (default: {DEFAULT_TEMPLATE.relative_to(PROJECT_ROOT)})")
    p.add_argument("--mode", choices=["fast", "hybrid", "claude"], default="fast",
                   help="Extraction backend (default: fast). 'hybrid' requires "
                        "opendataloader-pdf-hybrid running on 127.0.0.1:5002. "
                        "'claude' uses the Claude API (requires ANTHROPIC_API_KEY).")
    p.add_argument("--product", default="Operations Communication Manager",
                   help="Product name to write on the cover page")
    p.add_argument("--no-cache", action="store_true",
                   help="Re-run extraction even if cached JSON exists")
    p.add_argument("--json", dest="prebuilt_json", type=Path, default=None,
                   help="Use this pre-extracted opendataloader JSON instead of "
                        "running the extractor (skips Java requirement)")
    p.add_argument("--verbose", action="store_true", help="DEBUG logging")
    return p.parse_args()


def main() -> int:
    args = parse_args()
    setup_logging(args.verbose)

    if args.mode == "claude" and args.prebuilt_json is not None:
        logger.error("--json is not compatible with --mode claude")
        return 2

    try:
        mode = resolve_mode(args.mode)
        convert(
            pdf_path=args.pdf,
            template_path=args.template,
            output_path=args.output,
            mode=mode,
            product=args.product,
            use_cache=not args.no_cache,
            prebuilt_json=args.prebuilt_json,
        )
    except FileNotFoundError as e:
        logger.error("convert.failed reason=file_not_found %s", e)
        return 2
    except Exception as e:
        logger.exception("convert.failed reason=unexpected error=%s", e)
        return 1
    return 0


if __name__ == "__main__":
    sys.exit(main())
