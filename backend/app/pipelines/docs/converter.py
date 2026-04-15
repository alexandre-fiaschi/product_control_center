"""Docs pipeline — Block B: extract structured items + render CAE-templated DOCX.

Two public entry points, one per pipeline stage:

  - extract_release_notes(): downloaded → extracted (Claude API or cache hit)
  - render_release_notes(): extracted → converted (template + python-docx)

Each one is meant to be called from inside services.lifecycle.run_cell —
exceptions propagate out and run_cell records the failure on cell.last_run.

The two stages talk to each other via a persisted ReleaseNoteRecord JSON
on disk (record_json_path on the cell). That file is the editable "source
of truth" Unit 9's review view will let humans tweak between extract and
render.

PLAN_DOCS_PIPELINE.md §2 Block B / Unit 5.
"""

from __future__ import annotations

import hashlib
import json
import logging
from datetime import datetime, timezone
from pathlib import Path
from typing import Literal

from docx import Document
from docx.oxml.ns import qn
from docx.shared import Inches

from app.config import settings
from app.integrations.claude.client import ClaudeClient
from app.integrations.claude.extractor import (
    EXTRACTOR_VERSION as CURRENT_EXTRACTOR_VERSION,
    extract_release_note,
)
from app.integrations.pdf.image_extractor import extract_images
from app.state.models import PatchEntry
from app.state.release_notes_models import ReleaseNoteItem, ReleaseNoteRecord

logger = logging.getLogger("pipelines.docs.converter")


# ─────────────────────────── style constants ───────────────────────────

DEFAULT_BODY_STYLE = "Body Text"
DEFAULT_BULLET_STYLE = "Bullet 1"
DEFAULT_NUMBERED_STYLE = "Step 1"
TABLE_HEADING_STYLE = "Table Heading"
TABLE_TEXT_STYLE = "Table Text"

MAX_IMAGE_WIDTH = Inches(5.5)

COVER_INSTRUCTION_PHRASES = (
    "If this is your first ti",
    "Click on this text box and delete it",
    "Click this text box and delete it",
    "Update Customer Name",
)

TEMPLATE_BODY_START_INDEX = 38


# ─────────────────────────── cache helpers ───────────────────────────


def _cache_path_for(pdf_path: Path) -> Path:
    pdf_hash = hashlib.sha256(pdf_path.read_bytes()).hexdigest()
    return settings.docs_cache_dir / f"{pdf_hash}.json"


def _load_cached_record(pdf_path: Path) -> ReleaseNoteRecord | None:
    """Return the cached record for this PDF, or None on miss / version mismatch."""
    cache_file = _cache_path_for(pdf_path)
    if not cache_file.exists():
        return None
    try:
        record = ReleaseNoteRecord.model_validate_json(cache_file.read_text())
    except Exception as exc:
        logger.warning(
            "convert.extract.cache.invalid pdf=%s error=%s",
            pdf_path.name, exc,
        )
        return None
    if record.extractor_version != CURRENT_EXTRACTOR_VERSION:
        logger.info(
            "convert.extract.cache.stale pdf=%s cached_version=%d current_version=%d",
            pdf_path.name, record.extractor_version, CURRENT_EXTRACTOR_VERSION,
        )
        return None
    logger.info(
        "convert.extract.cache.hit pdf=%s cache=%s",
        pdf_path.name, cache_file.name[:12],
    )
    return record


def _save_cached_record(pdf_path: Path, record: ReleaseNoteRecord) -> None:
    settings.docs_cache_dir.mkdir(parents=True, exist_ok=True)
    cache_file = _cache_path_for(pdf_path)
    cache_file.write_text(record.model_dump_json(indent=2))
    logger.info(
        "convert.extract.cache.saved pdf=%s cache=%s",
        pdf_path.name, cache_file.name[:12],
    )


# ─────────────────────────── extract stage ───────────────────────────


def extract_release_notes(
    patch: PatchEntry,
    *,
    product_id: str,
    version: str,
    claude_client: ClaudeClient | None,
) -> Literal["extracted", "skipped_no_api"]:
    """Extract structured items from the downloaded source PDF.

    Workflow status transitions (only this function may touch them):
      - Cache hit OR fresh API call succeeds: downloaded → extracted
      - Cache miss + claude_client is None: status untouched (clean skip)
      - Exception: WORKFLOW STATUS UNTOUCHED. run_cell records the failure.

    See PLAN_DOCS_PIPELINE.md §2 Block B / Unit 5 for the contract.
    """
    cell = patch.release_notes
    if cell.source_pdf_path is None:
        raise ValueError(
            f"extract_release_notes precondition failed: source_pdf_path is None "
            f"(product={product_id} version={version})"
        )
    pdf_path = Path(cell.source_pdf_path).resolve()
    if not pdf_path.exists():
        raise FileNotFoundError(f"Source PDF missing: {pdf_path}")

    record = _load_cached_record(pdf_path)
    if record is None:
        if claude_client is None:
            logger.info(
                "convert.extract.skipped reason=claude_disabled product=%s version=%s",
                product_id, version,
            )
            return "skipped_no_api"
        logger.info(
            "convert.extract.start product=%s version=%s pdf=%s",
            product_id, version, pdf_path.name,
        )
        manifest = extract_images(pdf_path)
        record = extract_release_note(
            pdf_path,
            manifest,
            version=version,
            claude_client=claude_client,
        )
        _save_cached_record(pdf_path, record)

    # Persist the record JSON next to the source PDF so render (and Unit 9)
    # can read it directly without going through the cache directory.
    record_json_path = pdf_path.parent / f"{version}.json"
    record_json_path.write_text(record.model_dump_json(indent=2))

    cell.status = "extracted"
    cell.extracted_at = datetime.now(timezone.utc)
    cell.record_json_path = str(record_json_path)
    logger.info(
        "convert.extract.success product=%s version=%s items=%d record=%s",
        product_id, version, len(record.items), record_json_path,
    )
    return "extracted"


# ─────────────────────────── render stage ───────────────────────────


def render_release_notes(
    patch: PatchEntry,
    *,
    product_id: str,
    version: str,
    template_path: Path,
) -> None:
    """Render a CAE-templated DOCX from a previously-extracted record JSON.

    Workflow status transitions:
      - Success: extracted → converted (DOCX file written, generated_docx_path set)
      - Exception: WORKFLOW STATUS UNTOUCHED. run_cell records the failure.

    Idempotent: re-running on a 'converted' cell would overwrite the DOCX,
    but the orchestrator filters by status == "extracted" so that path only
    runs from Unit 9's manual re-render endpoint with a relaxed precondition.
    """
    cell = patch.release_notes
    if cell.record_json_path is None:
        raise ValueError(
            f"render_release_notes precondition failed: record_json_path is None "
            f"(product={product_id} version={version})"
        )
    record_path = Path(cell.record_json_path).resolve()
    if not record_path.exists():
        raise FileNotFoundError(f"Record JSON missing: {record_path}")
    if not template_path.exists():
        raise FileNotFoundError(f"Template not found: {template_path}")

    record = ReleaseNoteRecord.model_validate_json(record_path.read_text())

    # The image extractor wrote the PNGs to <pdf_dir>/images/. The PDF lives
    # in the same parent folder as the record JSON.
    images_dir = record_path.parent / "images"

    logger.info(
        "convert.render.start product=%s version=%s template=%s",
        product_id, version, template_path.name,
    )

    doc = Document(str(template_path))
    cover_count = patch_cover_page(doc, cover_replacements(product_id, version))
    logger.debug("convert.render.cover.patched runs=%d", cover_count)

    cleaned = clean_cover_textboxes(doc)
    logger.debug("convert.render.cover.textboxes.cleaned paragraphs=%d", cleaned)

    stripped_paras, stripped_tables = strip_template_body(doc)
    logger.debug(
        "convert.render.template.body.stripped paragraphs=%d tables=%d",
        stripped_paras, stripped_tables,
    )

    if mark_toc_dirty(doc):
        logger.debug("convert.render.toc.dirty.marked status=ok")
    else:
        logger.warning("convert.render.toc.dirty.marked status=not_found")

    counts = _render_record(doc, record, images_dir)
    logger.info("convert.render.body.emitted breakdown=%s", counts)

    output_path = record_path.parent / f"{version}.docx"
    output_path.parent.mkdir(parents=True, exist_ok=True)
    doc.save(str(output_path))
    logger.info(
        "convert.render.success product=%s version=%s output=%s size=%d",
        product_id, version, output_path, output_path.stat().st_size,
    )

    cell.status = "converted"
    cell.converted_at = datetime.now(timezone.utc)
    cell.generated_docx_path = str(output_path)


# ─────────────────────────── render internals ───────────────────────────


def format_item_heading(item: ReleaseNoteItem) -> str:
    """Build the Heading 2 text: 'AM1393 [HAL] - Title'."""
    heading = item.am_card
    if item.customers:
        heading += f" [{', '.join(item.customers)}]"
    heading += f" - {item.title}"
    return heading


def _render_record(doc: Document, record: ReleaseNoteRecord, images_dir: Path) -> dict[str, int]:
    """Walk a ReleaseNoteRecord and emit content into the DOCX.

    Returns a counts dict for logging.
    """
    counts: dict[str, int] = {"sections": 0, "items": 0, "blocks": 0}

    sections: dict[str, list[ReleaseNoteItem]] = {}
    for item in record.items:
        sections.setdefault(item.section, []).append(item)

    for section_name, items in sections.items():
        add_styled_paragraph(doc, section_name, "Heading 1")
        counts["sections"] += 1

        for item in items:
            add_styled_paragraph(doc, format_item_heading(item), "Heading 2")
            counts["items"] += 1

            if item.summary:
                add_styled_paragraph(doc, item.summary, DEFAULT_BODY_STYLE)

            for block in item.body:
                counts["blocks"] += 1
                btype = block.type

                if btype == "paragraph":
                    add_styled_paragraph(doc, block.text, DEFAULT_BODY_STYLE)

                elif btype == "heading":
                    add_bold_body_paragraph(doc, block.text)

                elif btype == "image":
                    img_path = images_dir / f"{block.image_id}.png"
                    if img_path.exists():
                        try:
                            doc.add_picture(str(img_path), width=MAX_IMAGE_WIDTH)
                        except Exception as exc:
                            logger.warning(
                                "convert.render.image.embed.failed id=%s error=%s",
                                block.image_id, exc,
                            )
                    else:
                        logger.warning(
                            "convert.render.image.missing id=%s path=%s",
                            block.image_id, img_path,
                        )

                elif btype == "list":
                    style = DEFAULT_NUMBERED_STYLE if block.ordered else DEFAULT_BULLET_STYLE
                    for li_text in block.items:
                        add_styled_paragraph(doc, li_text, style)

                elif btype == "table":
                    if block.headers or block.rows:
                        n_cols = len(block.headers) if block.headers else (
                            max((len(r) for r in block.rows), default=0)
                        )
                        total_rows = (1 if block.headers else 0) + len(block.rows)
                        if n_cols > 0 and total_rows > 0:
                            table = doc.add_table(rows=total_rows, cols=n_cols)
                            row_idx = 0
                            if block.headers:
                                for c, h in enumerate(block.headers):
                                    cell = table.rows[0].cells[c]
                                    cell.paragraphs[0].text = h
                                    if style_safe(doc, TABLE_HEADING_STYLE):
                                        cell.paragraphs[0].style = doc.styles[TABLE_HEADING_STYLE]
                                row_idx = 1
                            for r in block.rows:
                                for c, val in enumerate(r):
                                    if c < n_cols:
                                        cell = table.rows[row_idx].cells[c]
                                        cell.paragraphs[0].text = val
                                        if style_safe(doc, TABLE_TEXT_STYLE):
                                            cell.paragraphs[0].style = doc.styles[TABLE_TEXT_STYLE]
                                row_idx += 1

                elif btype == "code":
                    add_styled_paragraph(doc, block.text, DEFAULT_BODY_STYLE)

    return counts


# ─────────────────────────── template helpers ───────────────────────────


def style_safe(doc: Document, name: str) -> str | None:
    """Return name if it exists in the doc's style table, else None."""
    try:
        doc.styles[name]
        return name
    except KeyError:
        return None


def add_styled_paragraph(doc: Document, text: str, style: str | None) -> None:
    if not text or not text.strip():
        return
    p = doc.add_paragraph(text)
    if style and style_safe(doc, style):
        p.style = doc.styles[style]


def add_bold_body_paragraph(doc: Document, text: str) -> None:
    """Body Text paragraph with the entire run bolded — for AM sub-headings
    that should appear inline rather than in the TOC."""
    if not text or not text.strip():
        return
    p = doc.add_paragraph()
    if style_safe(doc, DEFAULT_BODY_STYLE):
        p.style = doc.styles[DEFAULT_BODY_STYLE]
    run = p.add_run(text)
    run.bold = True


def cover_replacements(product: str, version: str) -> dict[str, str]:
    return {
        "External Business Document": "Release Notes",
        "Customer": "ACARS",
        "Sample Product Name": product,
        "Version #": f"Version {version}",
        "Date": datetime.now().strftime("%B %Y"),
    }


def patch_cover_page(doc: Document, replacements: dict[str, str]) -> int:
    """Walk every <w:t> element in the body (including text boxes) and
    replace by exact match. Returns number of runs replaced."""
    replaced = 0
    for t in doc.element.body.iter(qn("w:t")):
        if t.text in replacements:
            t.text = replacements[t.text]
            replaced += 1
    return replaced


def clean_cover_textboxes(doc: Document) -> int:
    """Wipe instructional text inside cover-page text boxes."""
    cleaned = 0
    for txbx in doc.element.body.iter(qn("w:txbxContent")):
        for p in txbx.iter(qn("w:p")):
            runs = list(p.iter(qn("w:t")))
            joined = "".join(r.text or "" for r in runs)
            if not joined.strip():
                continue
            if any(phrase in joined for phrase in COVER_INSTRUCTION_PHRASES):
                for r in runs:
                    r.text = ""
                cleaned += 1
    return cleaned


def mark_toc_dirty(doc: Document) -> bool:
    """Find the TOC field and mark it dirty so Word rebuilds it on open."""
    body = doc.element.body
    fldChar_tag = qn("w:fldChar")
    instrText_tag = qn("w:instrText")
    fldCharType_attr = qn("w:fldCharType")
    dirty_attr = qn("w:dirty")

    elements = list(body.iter())
    for i, elem in enumerate(elements):
        if elem.tag != fldChar_tag:
            continue
        if elem.get(fldCharType_attr) != "begin":
            continue
        for j in range(i + 1, min(i + 6, len(elements))):
            nxt = elements[j]
            if nxt.tag == instrText_tag and (nxt.text or "").lstrip().startswith("TOC"):
                elem.set(dirty_attr, "1")
                return True
            if nxt.tag == fldChar_tag:
                break
    return False


def strip_template_body(doc: Document) -> tuple[int, int]:
    """Remove the template's instructional example chapters (paragraphs
    after TEMPLATE_BODY_START_INDEX) while preserving cover page, TOC, and
    section breaks. Returns (removed_paras, removed_tables)."""
    body = doc.element.body
    p_tag = qn("w:p")
    tbl_tag = qn("w:tbl")
    sectPr_tag = qn("w:sectPr")

    para_index = 0
    removed_paras = 0
    removed_tables = 0

    for child in list(body):
        tag = child.tag
        if tag == p_tag:
            current_index = para_index
            para_index += 1
            if current_index < TEMPLATE_BODY_START_INDEX:
                continue
            pPr = child.find(qn("w:pPr"))
            if pPr is not None and pPr.find(qn("w:sectPr")) is not None:
                continue
            body.remove(child)
            removed_paras += 1
        elif tag == tbl_tag:
            if para_index < TEMPLATE_BODY_START_INDEX:
                continue
            body.remove(child)
            removed_tables += 1
        elif tag == sectPr_tag:
            continue

    return removed_paras, removed_tables
